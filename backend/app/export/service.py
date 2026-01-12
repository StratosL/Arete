import io
import json
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

from docx import Document
from docx.shared import Inches

from app.core.database import get_supabase_service_client
from app.core.llm import get_llm_response


class ExportService:
    """Handle resume export to PDF and DOCX formats"""

    # Comprehensive skill categorization - expanded and validated
    # Rule: Each skill appears in EXACTLY one category
    KNOWN_SKILLS = {
        'Languages': {
            # Core programming languages
            'python', 'javascript', 'typescript', 'java', 'c++', 'c#', 'c',
            'go', 'golang', 'rust', 'ruby', 'php', 'swift', 'kotlin', 'scala',
            'r', 'perl', 'lua', 'haskell', 'elixir', 'clojure', 'erlang',
            'objective-c', 'dart', 'groovy', 'f#', 'ocaml', 'julia',
            # Query/markup languages (programming-adjacent)
            'sql', 'plsql', 'tsql', 'nosql', 'graphql',
            'html', 'css', 'scss', 'sass', 'less',
            # Shell scripting
            'bash', 'shell', 'powershell', 'zsh',
        },
        'Frontend': {
            # JavaScript frameworks/libraries
            'react', 'react.js', 'reactjs', 'vue', 'vue.js', 'vuejs',
            'angular', 'angularjs', 'svelte', 'solid', 'solidjs',
            'next.js', 'nextjs', 'nuxt', 'nuxt.js', 'nuxtjs',
            'gatsby', 'remix', 'astro', 'qwik',
            # State management
            'redux', 'mobx', 'zustand', 'recoil', 'pinia', 'vuex',
            # CSS frameworks
            'tailwind', 'tailwindcss', 'bootstrap', 'material-ui', 'mui',
            'chakra', 'chakra-ui', 'styled-components', 'emotion',
            'ant design', 'antd', 'bulma', 'foundation',
            # Build tools (frontend-specific)
            'webpack', 'vite', 'parcel', 'rollup', 'esbuild', 'turbopack',
            # Other frontend
            'jquery', 'backbone', 'ember', 'alpine.js', 'htmx',
            'storybook', 'chromatic',
        },
        'Backend': {
            # Node.js ecosystem
            'node.js', 'nodejs', 'node', 'express', 'express.js', 'expressjs',
            'nestjs', 'nest.js', 'fastify', 'koa', 'hapi',
            # Python frameworks
            'django', 'flask', 'fastapi', 'tornado', 'pyramid', 'bottle',
            'celery', 'dramatiq',
            # Java/JVM
            'spring', 'spring boot', 'springboot', 'quarkus', 'micronaut',
            'play framework', 'dropwizard', 'vert.x',
            # Ruby
            'rails', 'ruby on rails', 'sinatra', 'hanami',
            # PHP
            'laravel', 'symfony', 'codeigniter', 'yii', 'cakephp', 'slim',
            # .NET
            'asp.net', 'asp.net core', '.net', '.net core', 'blazor',
            # Go
            'gin', 'echo', 'fiber', 'chi', 'gorilla',
            # Rust
            'actix', 'actix-web', 'rocket', 'axum', 'warp',
            # API technologies
            'rest', 'restful', 'rest api', 'grpc', 'soap', 'websocket',
            'openapi', 'swagger',
        },
        'Databases': {
            # Relational databases
            'postgresql', 'postgres', 'mysql', 'mariadb', 'sqlite',
            'oracle', 'sql server', 'mssql', 'db2', 'cockroachdb',
            # NoSQL databases
            'mongodb', 'mongo', 'couchdb', 'couchbase', 'cassandra',
            'dynamodb', 'cosmosdb', 'firestore', 'fauna', 'faunadb',
            # Key-value/cache
            'redis', 'memcached', 'etcd', 'hazelcast',
            # Search engines
            'elasticsearch', 'elastic', 'opensearch', 'solr', 'algolia',
            'meilisearch', 'typesense',
            # Graph databases
            'neo4j', 'arangodb', 'dgraph', 'tigergraph',
            # Time-series
            'influxdb', 'timescaledb', 'prometheus', 'clickhouse',
            # ORMs and query builders
            'prisma', 'sequelize', 'typeorm', 'sqlalchemy', 'knex',
            'drizzle', 'mongoose', 'hibernate', 'entity framework',
            # BaaS with DB
            'supabase', 'firebase', 'appwrite', 'pocketbase', 'neon',
            'planetscale', 'turso',
        },
        'Cloud & DevOps': {
            # Cloud providers
            'aws', 'amazon web services', 'gcp', 'google cloud',
            'azure', 'microsoft azure', 'digitalocean', 'linode',
            'heroku', 'vercel', 'netlify', 'railway', 'render', 'fly.io',
            # Containers
            'docker', 'podman', 'containerd', 'buildah',
            # Orchestration
            'kubernetes', 'k8s', 'openshift', 'rancher', 'nomad',
            'docker swarm', 'docker compose', 'helm',
            # CI/CD
            'ci/cd', 'cicd', 'github actions', 'gitlab ci', 'jenkins',
            'circleci', 'travis ci', 'azure devops', 'bitbucket pipelines',
            'argo cd', 'argocd', 'flux', 'tekton', 'drone',
            # Infrastructure as Code
            'terraform', 'pulumi', 'cloudformation', 'ansible',
            'chef', 'puppet', 'saltstack', 'crossplane',
            # Monitoring & observability
            'grafana', 'datadog', 'new relic', 'splunk', 'dynatrace',
            'elk', 'elk stack', 'loki', 'jaeger', 'zipkin', 'opentelemetry',
            # Service mesh & networking
            'istio', 'envoy', 'linkerd', 'consul', 'nginx', 'haproxy',
            'traefik', 'kong', 'cloudflare',
            # OS & infrastructure
            'linux', 'ubuntu', 'centos', 'debian', 'rhel', 'alpine',
            # Message queues
            'kafka', 'rabbitmq', 'sqs', 'sns', 'pubsub', 'nats',
            'activemq', 'zeromq',
            # Serverless
            'lambda', 'aws lambda', 'cloud functions', 'azure functions',
            'cloudflare workers', 'edge functions',
        },
        'Tools': {
            # Version control
            'git', 'github', 'gitlab', 'bitbucket', 'svn', 'mercurial',
            # IDEs & editors
            'vs code', 'vscode', 'visual studio', 'intellij', 'pycharm',
            'webstorm', 'rider', 'vim', 'neovim', 'emacs', 'sublime',
            # Testing frameworks
            'jest', 'mocha', 'jasmine', 'vitest', 'playwright', 'cypress',
            'selenium', 'puppeteer', 'pytest', 'unittest', 'rspec',
            'junit', 'testng', 'xunit', 'nunit',
            # API testing
            'postman', 'insomnia', 'httpie', 'curl', 'bruno',
            # Package managers
            'npm', 'yarn', 'pnpm', 'pip', 'poetry', 'conda',
            'maven', 'gradle', 'cargo', 'gem', 'composer',
            # Project management
            'jira', 'trello', 'asana', 'linear', 'notion', 'confluence',
            'clickup', 'monday',
            # Design tools
            'figma', 'sketch', 'adobe xd', 'invision', 'zeplin',
            # Documentation
            'markdown', 'latex', 'sphinx', 'docusaurus', 'mkdocs',
            # Linting & formatting
            'eslint', 'prettier', 'black', 'flake8', 'mypy', 'ruff',
            'rubocop', 'checkstyle', 'spotless',
            # Code quality
            'sonarqube', 'codeclimate', 'codecov', 'coveralls',
            # AI/ML tools (development tools, not frameworks)
            'jupyter', 'jupyter notebook', 'colab', 'kaggle',
        },
    }

    # Skill name normalization (canonical forms)
    SKILL_ALIASES = {
        # JavaScript ecosystem
        'js': 'JavaScript', 'javascript': 'JavaScript',
        'ts': 'TypeScript', 'typescript': 'TypeScript',
        'react.js': 'React', 'reactjs': 'React', 'react': 'React',
        'vue.js': 'Vue', 'vuejs': 'Vue', 'vue': 'Vue',
        'angular.js': 'Angular', 'angularjs': 'Angular',
        'node.js': 'Node.js', 'nodejs': 'Node.js', 'node': 'Node.js',
        'next.js': 'Next.js', 'nextjs': 'Next.js',
        'express.js': 'Express', 'expressjs': 'Express',
        'nest.js': 'NestJS', 'nestjs': 'NestJS',
        # Databases
        'postgres': 'PostgreSQL', 'postgresql': 'PostgreSQL',
        'mongo': 'MongoDB', 'mongodb': 'MongoDB',
        'mssql': 'SQL Server', 'sql server': 'SQL Server',
        # Cloud & DevOps
        'k8s': 'Kubernetes', 'kubernetes': 'Kubernetes',
        'aws': 'AWS', 'amazon web services': 'AWS',
        'gcp': 'GCP', 'google cloud': 'GCP', 'google cloud platform': 'GCP',
        'azure': 'Azure', 'microsoft azure': 'Azure',
        'cicd': 'CI/CD', 'ci/cd': 'CI/CD',
        # Frameworks
        'rails': 'Ruby on Rails', 'ruby on rails': 'Ruby on Rails',
        'spring boot': 'Spring Boot', 'springboot': 'Spring Boot',
        'asp.net core': 'ASP.NET Core', 'aspnet': 'ASP.NET',
        'fastapi': 'FastAPI',
        # Tools
        'vscode': 'VS Code', 'vs code': 'VS Code',
        'tailwindcss': 'Tailwind CSS', 'tailwind': 'Tailwind CSS',
        # CSS
        'scss': 'SCSS/Sass', 'sass': 'SCSS/Sass',
    }

    # Skills that should NEVER be in certain categories (validation rules)
    CATEGORY_EXCLUSIONS = {
        'Languages': {'react', 'vue', 'angular', 'django', 'flask', 'express',
                      'postgresql', 'mongodb', 'redis', 'docker', 'aws'},
        'Frontend': {'python', 'java', 'go', 'rust', 'postgresql', 'mongodb',
                     'docker', 'kubernetes', 'aws'},
        'Backend': {'html', 'css', 'tailwind', 'react', 'vue', 'angular'},
        'Databases': {'python', 'javascript', 'react', 'docker', 'aws'},
        'Cloud & DevOps': {'react', 'vue', 'angular', 'python', 'javascript'},
    }

    CATEGORIZATION_PROMPT = """You are a technical skills categorization expert. Categorize each skill precisely.

STRICT RULES - Follow these exactly:

**Languages** (Programming/scripting languages ONLY):
- Python, JavaScript, TypeScript, Java, C++, C#, Go, Rust, Ruby, PHP, Swift, Kotlin
- SQL, GraphQL (query languages)
- HTML, CSS (markup/styling - they ARE languages)
- Bash, PowerShell (shell scripting)
- NEVER include frameworks here (React is NOT a language, it uses JavaScript)

**Frontend** (Browser/UI technologies):
- Frameworks: React, Vue, Angular, Svelte, Next.js, Nuxt
- State: Redux, MobX, Zustand, Pinia
- Styling: Tailwind, Bootstrap, Material-UI, Chakra
- Build: Webpack, Vite, Parcel

**Backend** (Server-side frameworks/runtimes):
- Node.js/Express/NestJS/Fastify
- Python: Django, Flask, FastAPI
- Java: Spring Boot, Quarkus
- Ruby on Rails, Laravel, ASP.NET
- API: REST, gRPC, GraphQL (as protocol)

**Databases** (Data storage/ORMs):
- SQL: PostgreSQL, MySQL, SQLite, SQL Server
- NoSQL: MongoDB, Redis, Cassandra, DynamoDB
- Search: Elasticsearch, Algolia
- ORMs: Prisma, Sequelize, SQLAlchemy, TypeORM
- BaaS: Supabase, Firebase

**Cloud & DevOps** (Infrastructure/deployment):
- Cloud: AWS, GCP, Azure, Vercel, Netlify
- Containers: Docker, Kubernetes, Helm
- CI/CD: GitHub Actions, Jenkins, GitLab CI
- IaC: Terraform, Ansible, Pulumi
- Monitoring: Grafana, Prometheus, Datadog

**Tools** (Development tools):
- VCS: Git, GitHub, GitLab
- Testing: Jest, Pytest, Cypress, Playwright
- IDEs: VS Code, IntelliJ
- Package managers: npm, pip, Maven

**Other** (LAST RESORT - rarely used):
- Soft skills: Agile, Scrum, Leadership
- Non-tech domains only

Skills to categorize:
{skills}

Return ONLY valid JSON:
{{
  "Languages": [],
  "Frontend": [],
  "Backend": [],
  "Databases": [],
  "Cloud & DevOps": [],
  "Tools": [],
  "Other": []
}}"""

    def __init__(self):
        self.templates_dir = Path(__file__).parent / "templates"

    def _sort_projects(self, projects: list[dict]) -> list[dict]:
        """Sort projects: resume-sourced first, GitHub-sourced last.

        GitHub projects are identified by:
        - Having 'github_url' field populated
        - Having 'source' field set to 'github'
        - Having 'github.com' in any URL field
        """
        if not projects:
            return projects

        def is_github_project(project: dict) -> bool:
            # Check explicit github_url field
            if project.get('github_url'):
                return True
            # Check source field
            if project.get('source', '').lower() == 'github':
                return True
            # Check if any URL contains github.com
            for key in ['url', 'link', 'repo_url']:
                url = project.get(key, '')
                if url and 'github.com' in url.lower():
                    return True
            return False

        # Separate into resume projects and GitHub projects
        resume_projects = []
        github_projects = []

        for project in projects:
            if is_github_project(project):
                github_projects.append(project)
            else:
                resume_projects.append(project)

        # Return resume projects first, then GitHub projects
        return resume_projects + github_projects

    def _validate_skill_category(
        self, skill: str, assigned_category: str
    ) -> str | None:
        """Validate that a skill is correctly categorized.

        Returns the correct category if miscategorized, None if correct.
        """
        skill_lower = skill.lower().strip()

        # Check exclusion rules
        exclusions = self.CATEGORY_EXCLUSIONS.get(assigned_category, set())
        if skill_lower in exclusions:
            # Find the correct category
            for category, skills in self.KNOWN_SKILLS.items():
                if skill_lower in skills:
                    return category
            return 'Other'

        return None

    def _post_process_categorized_skills(
        self, categorized: dict[str, list[str]]
    ) -> dict[str, list[str]]:
        """Post-process categorized skills to fix common miscategorizations."""
        corrected = {cat: list(skills) for cat, skills in categorized.items()}

        # Validate each skill and move if needed
        moves = []  # (skill, from_category, to_category)

        for category, skills in corrected.items():
            for skill in skills:
                correct_cat = self._validate_skill_category(skill, category)
                if correct_cat:
                    moves.append((skill, category, correct_cat))

        # Apply moves
        for skill, from_cat, to_cat in moves:
            if skill in corrected.get(from_cat, []):
                corrected[from_cat].remove(skill)
            if to_cat not in corrected:
                corrected[to_cat] = []
            if skill not in corrected[to_cat]:
                corrected[to_cat].append(skill)

        # Remove empty categories
        return {cat: skills for cat, skills in corrected.items() if skills}

    def _normalize_skill(self, skill: str) -> str:
        """Normalize skill name to canonical form"""
        skill_lower = skill.lower().strip()
        return self.SKILL_ALIASES.get(skill_lower, skill.strip())
    
    def _quick_categorize(self, skill: str) -> str | None:
        """Try to categorize skill using known mappings. Returns None if unknown."""
        skill_lower = skill.lower().strip()
        
        # First pass: exact matches
        for category, skills in self.KNOWN_SKILLS.items():
            if skill_lower in skills:
                return category
        
        # Second pass: substring matches (for compound skills)
        for category, skills in self.KNOWN_SKILLS.items():
            if any(s in skill_lower for s in skills if len(s) > 2):  # Avoid single char matches
                return category
        
        return None
    
    async def _llm_categorize_skills(self, skills: list[str]) -> dict[str, list[str]]:
        """Use LLM to categorize unknown skills"""
        if not skills:
            return {}
        
        prompt = self.CATEGORIZATION_PROMPT.format(skills=json.dumps(skills))
        
        try:
            response = await get_llm_response([
                {"role": "user", "content": prompt}
            ])
            
            # Extract JSON from response (handle markdown code blocks)
            clean_response = response.strip()
            if "```json" in clean_response:
                clean_response = clean_response.split("```json")[1]
                clean_response = clean_response.split("```")[0]
            elif "```" in clean_response:
                clean_response = clean_response.split("```")[1]
                clean_response = clean_response.split("```")[0]
            
            json_start = clean_response.find('{')
            json_end = clean_response.rfind('}') + 1
            if json_start != -1 and json_end > json_start:
                raw_result = json.loads(clean_response[json_start:json_end])
                
                # Normalize category names from LLM response
                normalized = {}
                category_mapping = {
                    'languages': 'Languages',
                    'programming languages': 'Languages',
                    'frontend': 'Frontend',
                    'front-end': 'Frontend',
                    'front end': 'Frontend',
                    'backend': 'Backend',
                    'back-end': 'Backend',
                    'back end': 'Backend',
                    'databases': 'Databases',
                    'database': 'Databases',
                    'cloud & devops': 'Cloud & DevOps',
                    'cloud and devops': 'Cloud & DevOps',
                    'cloud': 'Cloud & DevOps',
                    'devops': 'Cloud & DevOps',
                    'infrastructure': 'Cloud & DevOps',
                    'tools': 'Tools',
                    'other': 'Other',
                    'soft skills': 'Other',
                    'methodologies': 'Other',
                }
                
                for key, value in raw_result.items():
                    if isinstance(value, list) and value:
                        norm_key = category_mapping.get(key.lower(), 'Other')
                        if norm_key not in normalized:
                            normalized[norm_key] = []
                        normalized[norm_key].extend(value)
                
                return normalized
                
        except Exception as e:
            print(f"LLM categorization failed: {e}")
        
        # Fallback: put all in Other
        return {"Other": skills}
    
    async def _deduplicate_and_categorize_skills(
        self, skills_dict: dict
    ) -> dict[str, list[str]]:
        """Deduplicate skills and categorize using quick-match + LLM fallback"""
        
        # Collect and deduplicate all skills
        all_skills = set()
        for category in ['technical', 'frameworks', 'tools', 'languages']:
            for skill in skills_dict.get(category, []):
                normalized = self._normalize_skill(skill)
                if normalized.lower() not in {s.lower() for s in all_skills}:
                    all_skills.add(normalized)
        
        # Categorize: quick-match first, collect unknowns for LLM
        categorized = {
            'Languages': [], 'Frontend': [], 'Backend': [],
            'Databases': [], 'Cloud & DevOps': [], 'Tools': [], 'Other': []
        }
        unknown_skills = []
        
        for skill in all_skills:
            category = self._quick_categorize(skill)
            if category:
                categorized[category].append(skill)
            else:
                unknown_skills.append(skill)
        
        # LLM categorization for unknown skills
        if unknown_skills:
            llm_results = await self._llm_categorize_skills(unknown_skills)
            for category, skills in llm_results.items():
                if category in categorized:
                    categorized[category].extend(skills)
                else:
                    categorized['Other'].extend(skills)

        # Post-process to fix any miscategorizations
        categorized = self._post_process_categorized_skills(categorized)

        # Sort each category and remove empty ones
        result = {}
        for category, skills in categorized.items():
            if skills:
                result[category] = sorted(set(skills), key=str.lower)

        return result
    
    async def export_resume(
        self, resume_id: str, format: str, template: str = "classic"
    ) -> tuple[bytes, str, str]:
        """Export resume in specified format with chosen template"""

        # Get resume data from database
        supabase = get_supabase_service_client()
        result = supabase.table("resumes").select("*").eq("id", resume_id).execute()

        if not result.data:
            raise ValueError(f"Resume {resume_id} not found")

        resume_record = result.data[0]
        # Use optimized data if available, otherwise fall back to parsed data
        resume_data = resume_record.get("optimized_data") or resume_record["parsed_data"]

        if format == "pdf":
            return await self._generate_pdf(resume_data, resume_id, template)
        elif format == "docx":
            return await self._generate_docx(resume_data, resume_id)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    async def _generate_pdf(
        self, resume_data: dict, resume_id: str, template: str = "classic"
    ) -> tuple[bytes, str, str]:
        """Generate PDF using selected template.

        - classic: ReportLab-generated PDF (ATS-optimized)
        - modern: HTML-based template (browser print to PDF)
        """

        personal = resume_data.get("personal_info", {})
        name = personal.get("name", "resume").replace(" ", "_")

        # Modern template returns HTML for browser-based PDF generation
        if template == "modern":
            html_content = await self._build_modern_html(resume_data)
            filename = f"{name}_resume_modern.html"
            return html_content.encode('utf-8'), "text/html", filename

        # Classic template uses ReportLab for direct PDF generation
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=12,
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=6,
            spaceBefore=12
        )
        
        skill_category_style = ParagraphStyle(
            'SkillCategory',
            parent=styles['Normal'],
            fontSize=10,
            spaceBefore=4,
            spaceAfter=2,
            leftIndent=0
        )
        
        story = []
        personal = resume_data["personal_info"]
        
        # Name
        story.append(Paragraph(personal["name"], title_style))
        
        # Contact info
        contact_info = []
        if personal.get("email"):
            contact_info.append(personal["email"])
        if personal.get("phone"):
            contact_info.append(personal["phone"])
        if personal.get("location"):
            contact_info.append(personal["location"])
        
        if contact_info:
            story.append(Paragraph(" | ".join(contact_info), styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Experience
        if resume_data.get("experience"):
            story.append(Paragraph("Experience", heading_style))
            for exp in resume_data["experience"]:
                title_company = f"<b>{exp['title']} - {exp['company']}</b>"
                story.append(Paragraph(title_company, styles['Normal']))
                story.append(Paragraph(f"<i>{exp['duration']}</i>", styles['Normal']))
                for desc in exp["description"]:
                    story.append(Paragraph(f"• {desc}", styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Skills - deduplicated and categorized via LLM
        if resume_data.get("skills"):
            story.append(Paragraph("Skills", heading_style))
            skills_data = await self._deduplicate_and_categorize_skills(
                resume_data["skills"]
            )
            
            # Define display order for categories
            category_order = [
                'Languages', 'Frontend', 'Backend', 'Databases',
                'Cloud & DevOps', 'Tools', 'Other'
            ]
            
            for category in category_order:
                if category in skills_data and skills_data[category]:
                    skills_list = skills_data[category]
                    skills_line = f"<b>{category}:</b> {', '.join(skills_list)}"
                    story.append(Paragraph(skills_line, skill_category_style))
            
            story.append(Spacer(1, 6))
        
        # Projects (sorted: resume projects first, GitHub-sourced last)
        if resume_data.get("projects"):
            story.append(Paragraph("Projects", heading_style))
            sorted_projects = self._sort_projects(resume_data["projects"])
            for project in sorted_projects:
                story.append(Paragraph(f"<b>{project['name']}</b>", styles['Normal']))
                story.append(Paragraph(project['description'], styles['Normal']))
                if project.get("technologies"):
                    tech_list = f"<b>Technologies:</b> {', '.join(project['technologies'])}"
                    story.append(Paragraph(tech_list, styles['Normal']))
                story.append(Spacer(1, 6))

        # Education
        if resume_data.get("education"):
            story.append(Paragraph("Education", heading_style))
            for edu in resume_data["education"]:
                story.append(Paragraph(f"<b>{edu['degree']}</b>", styles['Normal']))
                edu_info = edu['institution']
                if edu.get("graduation_year"):
                    edu_info += f" | {edu['graduation_year']}"
                if edu.get("gpa"):
                    edu_info += f" | GPA: {edu['gpa']}"
                story.append(Paragraph(edu_info, styles['Normal']))
        
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()

        filename = f"{personal['name'].replace(' ', '_')}_resume.pdf"
        return pdf_bytes, "application/pdf", filename
    
    async def _generate_docx(self, resume_data: dict, resume_id: str) -> tuple[bytes, str, str]:
        """Generate ATS-compliant DOCX using python-docx"""
        
        doc = Document()
        
        # Personal Info
        personal = resume_data["personal_info"]
        doc.add_heading(personal["name"], 0)
        
        contact_info = []
        if personal.get("email"):
            contact_info.append(personal["email"])
        if personal.get("phone"):
            contact_info.append(personal["phone"])
        if personal.get("location"):
            contact_info.append(personal["location"])
        
        if contact_info:
            doc.add_paragraph(" | ".join(contact_info))
        
        # Experience
        if resume_data.get("experience"):
            doc.add_heading("Experience", 1)
            for exp in resume_data["experience"]:
                p = doc.add_paragraph()
                p.add_run(f"{exp['title']} - {exp['company']}").bold = True
                doc.add_paragraph(exp["duration"])
                for desc in exp["description"]:
                    doc.add_paragraph(f"• {desc}")
        
        # Skills - deduplicated and categorized via LLM
        if resume_data.get("skills"):
            doc.add_heading("Skills", 1)
            skills_data = await self._deduplicate_and_categorize_skills(
                resume_data["skills"]
            )
            
            for category, skills_list in skills_data.items():
                if skills_list:
                    p = doc.add_paragraph()
                    p.add_run(f"{category}: ").bold = True
                    p.add_run(', '.join(skills_list))
        
        # Projects (sorted: resume projects first, GitHub-sourced last)
        if resume_data.get("projects"):
            doc.add_heading("Projects", 1)
            sorted_projects = self._sort_projects(resume_data["projects"])
            for project in sorted_projects:
                p = doc.add_paragraph()
                p.add_run(project["name"]).bold = True
                doc.add_paragraph(project["description"])
                if project.get("technologies"):
                    doc.add_paragraph(f"Technologies: {', '.join(project['technologies'])}")

        # Education
        if resume_data.get("education"):
            doc.add_heading("Education", 1)
            for edu in resume_data["education"]:
                if edu.get("degree") and edu.get("institution"):
                    doc.add_paragraph(f"{edu['degree']} - {edu['institution']}")
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"{personal['name'].replace(' ', '_')}_resume.docx"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return buffer.getvalue(), content_type, filename
    
    def _build_html(self, resume_data: dict) -> str:
        """Build HTML content for PDF generation with inline CSS"""
        
        personal = resume_data["personal_info"]
        
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{personal['name']} - Resume</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    font-size: 11pt;
                    line-height: 1.4;
                    margin: 0.5in;
                    color: #000;
                }}
                .resume {{
                    max-width: 8.5in;
                }}
                header {{
                    text-align: center;
                    margin-bottom: 20pt;
                }}
                h1 {{
                    font-size: 18pt;
                    font-weight: bold;
                    margin: 0 0 8pt 0;
                }}
                h2 {{
                    font-size: 14pt;
                    font-weight: bold;
                    margin: 16pt 0 8pt 0;
                    border-bottom: 1pt solid #000;
                }}
                h3 {{
                    font-size: 12pt;
                    font-weight: bold;
                    margin: 12pt 0 4pt 0;
                }}
                .contact {{
                    font-size: 10pt;
                    margin-bottom: 8pt;
                }}
                .duration {{
                    font-style: italic;
                    margin: 0 0 8pt 0;
                }}
                ul {{
                    margin: 8pt 0;
                    padding-left: 20pt;
                }}
                li {{
                    margin-bottom: 4pt;
                }}
                section {{
                    margin-bottom: 16pt;
                }}
                .job, .project {{
                    margin-bottom: 12pt;
                }}
            </style>
        </head>
        <body>
            <div class="resume">
                <header>
                    <h1>{personal['name']}</h1>
                    <div class="contact">
        """
        
        contact_info = []
        if personal.get("email"):
            contact_info.append(personal["email"])
        if personal.get("phone"):
            contact_info.append(personal["phone"])
        if personal.get("location"):
            contact_info.append(personal["location"])
        
        html += " | ".join(contact_info)
        html += """
                    </div>
                </header>
        """
        
        # Experience
        if resume_data.get("experience"):
            html += "<section><h2>Experience</h2>"
            for exp in resume_data["experience"]:
                html += f"""
                <div class="job">
                    <h3>{exp['title']} - {exp['company']}</h3>
                    <p class="duration">{exp['duration']}</p>
                    <ul>
                """
                for desc in exp["description"]:
                    html += f"<li>{desc}</li>"
                html += "</ul></div>"
            html += "</section>"
        
        # Skills
        if resume_data.get("skills"):
            html += "<section><h2>Skills</h2>"
            skills = resume_data["skills"]
            if skills.get("technical"):
                html += f"<p><strong>Technical:</strong> {', '.join(skills['technical'])}</p>"
            if skills.get("frameworks"):
                html += f"<p><strong>Frameworks:</strong> {', '.join(skills['frameworks'])}</p>"
            html += "</section>"
        
        # Projects (sorted: resume projects first, GitHub-sourced last)
        if resume_data.get("projects"):
            html += "<section><h2>Projects</h2>"
            sorted_projects = self._sort_projects(resume_data["projects"])
            for project in sorted_projects:
                html += f"""
                <div class="project">
                    <h3>{project['name']}</h3>
                    <p>{project['description']}</p>
                """
                if project.get("technologies"):
                    tech_list = ', '.join(project['technologies'])
                    tech_str = f"<strong>Technologies:</strong> {tech_list}"
                    html += f"<p>{tech_str}</p>"
                html += "</div>"
            html += "</section>"

        # Education
        if resume_data.get("education"):
            html += "<section><h2>Education</h2>"
            for edu in resume_data["education"]:
                if edu.get("degree") and edu.get("institution"):
                    html += f"<p>{edu['degree']} - {edu['institution']}"
                    if edu.get("graduation_year"):
                        html += f" ({edu['graduation_year']})"
                    html += "</p>"
            html += "</section>"
        
        html += """
            </div>
        </body>
        </html>
        """
        
        return html
    
    def _build_printable_html(self, resume_data: dict) -> str:
        """Build HTML optimized for browser PDF printing"""
        
        personal = resume_data["personal_info"]
        
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{personal['name']} - Resume</title>
            <style>
                @media print {{
                    body {{ margin: 0.5in; }}
                    .no-print {{ display: none; }}
                }}
                body {{ 
                    font-family: Arial, sans-serif; 
                    font-size: 11pt; 
                    line-height: 1.4;
                    margin: 0.5in;
                    color: #000;
                }}
                h1 {{ 
                    font-size: 18pt; 
                    margin-bottom: 10pt; 
                    border-bottom: 2pt solid #000;
                    padding-bottom: 5pt;
                }}
                h2 {{ 
                    font-size: 14pt; 
                    margin: 15pt 0 8pt 0; 
                    color: #333;
                }}
                h3 {{ 
                    font-size: 12pt; 
                    margin: 10pt 0 5pt 0; 
                    font-weight: bold;
                }}
                p {{ margin: 5pt 0; }}
                ul {{ margin: 5pt 0; padding-left: 20pt; }}
                li {{ margin-bottom: 3pt; }}
                .contact {{ 
                    font-size: 10pt; 
                    margin-bottom: 15pt; 
                    color: #666;
                }}
                .duration {{ 
                    font-style: italic; 
                    color: #666; 
                    margin-bottom: 5pt;
                }}
                .print-instruction {{
                    background: #f0f0f0;
                    padding: 10pt;
                    margin-bottom: 20pt;
                    border-left: 4pt solid #007acc;
                }}
            </style>
        </head>
        <body>
            <div class="print-instruction no-print">
                <strong>To save as PDF:</strong> Use your browser's Print function (Ctrl+P) 
                and select "Save as PDF"
            </div>
            
            <h1>{personal['name']}</h1>
            <div class="contact">
        """
        
        contact_info = []
        if personal.get("email"):
            contact_info.append(personal["email"])
        if personal.get("phone"):
            contact_info.append(personal["phone"])
        if personal.get("location"):
            contact_info.append(personal["location"])
        if personal.get("github"):
            contact_info.append(f"GitHub: {personal['github']}")
        if personal.get("linkedin"):
            contact_info.append(f"LinkedIn: {personal['linkedin']}")
        
        html += " | ".join(contact_info) + "</div>"
        
        # Experience
        if resume_data.get("experience"):
            html += "<h2>Experience</h2>"
            for exp in resume_data["experience"]:
                html += f"<h3>{exp['title']} - {exp['company']}</h3>"
                html += f"<div class='duration'>{exp['duration']}</div>"
                html += "<ul>"
                for desc in exp["description"]:
                    html += f"<li>{desc}</li>"
                html += "</ul>"
        
        # Skills
        if resume_data.get("skills"):
            html += "<h2>Skills</h2>"
            skills = resume_data["skills"]
            if skills.get("technical"):
                html += f"<p><strong>Technical:</strong> {', '.join(skills['technical'])}</p>"
            if skills.get("frameworks"):
                html += f"<p><strong>Frameworks:</strong> {', '.join(skills['frameworks'])}</p>"
            if skills.get("tools"):
                html += f"<p><strong>Tools:</strong> {', '.join(skills['tools'])}</p>"

        # Projects (sorted: resume projects first, GitHub-sourced last)
        if resume_data.get("projects"):
            html += "<h2>Projects</h2>"
            sorted_projects = self._sort_projects(resume_data["projects"])
            for project in sorted_projects:
                html += f"<h3>{project['name']}</h3>"
                html += f"<p>{project['description']}</p>"
                if project.get("technologies"):
                    tech_list = ', '.join(project['technologies'])
                    html += f"<p><strong>Technologies:</strong> {tech_list}</p>"

        # Education
        if resume_data.get("education"):
            html += "<h2>Education</h2>"
            for edu in resume_data["education"]:
                html += f"<h3>{edu['degree']}</h3>"
                html += f"<p>{edu['institution']}"
                if edu.get("graduation_year"):
                    html += f" | {edu['graduation_year']}"
                if edu.get("gpa"):
                    html += f" | GPA: {edu['gpa']}"
                html += "</p>"
        
        html += """
        </body>
        </html>
        """
        
        return html

    def _build_simple_html(self, resume_data: dict) -> str:
        """Build minimal HTML content as fallback for WeasyPrint issues"""
        
        personal = resume_data["personal_info"]
        
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{personal['name']} - Resume</title>
            <style>
                body {{ font-family: Arial, sans-serif; font-size: 11pt; margin: 0.5in; }}
                h1 {{ font-size: 18pt; margin-bottom: 10pt; }}
                h2 {{ font-size: 14pt; margin: 15pt 0 8pt 0; }}
                h3 {{ font-size: 12pt; margin: 10pt 0 5pt 0; }}
                p {{ margin: 5pt 0; }}
                ul {{ margin: 5pt 0; padding-left: 20pt; }}
            </style>
        </head>
        <body>
            <h1>{personal['name']}</h1>
            <p>
        """
        
        contact_info = []
        if personal.get("email"):
            contact_info.append(personal["email"])
        if personal.get("phone"):
            contact_info.append(personal["phone"])
        if personal.get("location"):
            contact_info.append(personal["location"])
        
        html += " | ".join(contact_info) + "</p>"
        
        # Experience
        if resume_data.get("experience"):
            html += "<h2>Experience</h2>"
            for exp in resume_data["experience"]:
                html += f"<h3>{exp['title']} - {exp['company']}</h3>"
                html += f"<p><em>{exp['duration']}</em></p>"
                html += "<ul>"
                for desc in exp["description"]:
                    html += f"<li>{desc}</li>"
                html += "</ul>"
        
        # Skills
        if resume_data.get("skills"):
            html += "<h2>Skills</h2>"
            skills = resume_data["skills"]
            if skills.get("technical"):
                html += f"<p><strong>Technical:</strong> {', '.join(skills['technical'])}</p>"
            if skills.get("frameworks"):
                html += f"<p><strong>Frameworks:</strong> {', '.join(skills['frameworks'])}</p>"

        # Projects (sorted: resume projects first, GitHub-sourced last)
        if resume_data.get("projects"):
            html += "<h2>Projects</h2>"
            sorted_projects = self._sort_projects(resume_data["projects"])
            for project in sorted_projects:
                html += f"<h3>{project['name']}</h3>"
                html += f"<p>{project['description']}</p>"
                if project.get("technologies"):
                    tech_list = ', '.join(project['technologies'])
                    html += f"<p><strong>Technologies:</strong> {tech_list}</p>"

        # Education
        if resume_data.get("education"):
            html += "<h2>Education</h2>"
            for edu in resume_data["education"]:
                if edu.get("degree") and edu.get("institution"):
                    html += f"<p>{edu['degree']} - {edu['institution']}"
                    if edu.get("graduation_year"):
                        html += f" ({edu['graduation_year']})"
                    html += "</p>"

        html += "</body></html>"
        return html

    def _load_template_css(self, template_name: str) -> str:
        """Load CSS from template file"""
        css_file = self.templates_dir / f"{template_name}.css"
        if css_file.exists():
            return css_file.read_text()
        return ""

    async def _build_modern_html(self, resume_data: dict) -> str:
        """Build HTML content using modern template with CSS styling"""

        personal = resume_data.get("personal_info", {})
        css_content = self._load_template_css("modern")

        # Build contact items
        contact_items = []
        if personal.get("email"):
            email = personal["email"]
            contact_items.append(
                f'<span class="contact-item">'
                f'<a href="mailto:{email}">{email}</a></span>'
            )
        if personal.get("phone"):
            contact_items.append(
                f'<span class="contact-item">{personal["phone"]}</span>'
            )
        if personal.get("location"):
            contact_items.append(
                f'<span class="contact-item">{personal["location"]}</span>'
            )
        if personal.get("linkedin"):
            linkedin = personal["linkedin"]
            display_linkedin = linkedin.replace("https://", "").replace("http://", "")
            contact_items.append(
                f'<span class="contact-item">'
                f'<a href="{linkedin}">{display_linkedin}</a></span>'
            )
        if personal.get("github"):
            github = personal["github"]
            display_github = github.replace("https://", "").replace("http://", "")
            contact_items.append(
                f'<span class="contact-item">'
                f'<a href="{github}">{display_github}</a></span>'
            )

        contact_html = " | ".join(contact_items)

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{personal.get('name', 'Resume')} - Resume</title>
    <style>
{css_content}
    </style>
</head>
<body>
    <div class="resume-container">
        <header class="resume-header">
            <h1 class="resume-name">{personal.get('name', '')}</h1>
            <div class="resume-contact">
                {contact_html}
            </div>
        </header>
"""

        # Summary section (if present)
        if resume_data.get("summary"):
            html += f"""
        <section class="resume-section">
            <h2 class="section-title">Summary</h2>
            <p class="resume-summary">{resume_data['summary']}</p>
        </section>
"""

        # Experience section
        if resume_data.get("experience"):
            html += """
        <section class="resume-section">
            <h2 class="section-title">Experience</h2>
"""
            for exp in resume_data["experience"]:
                title = exp.get('title', '')
                company = exp.get('company', '')
                duration = exp.get('duration', '')

                html += f"""
            <div class="experience-item">
                <div class="experience-header">
                    <div>
                        <span class="experience-title">{title}</span>
                        <span class="experience-company"> - {company}</span>
                    </div>
                    <span class="experience-duration">{duration}</span>
                </div>
                <ul class="experience-description">
"""
                for desc in exp.get("description", []):
                    # Escape HTML entities
                    safe_desc = (desc.replace("&", "&amp;")
                                    .replace("<", "&lt;")
                                    .replace(">", "&gt;"))
                    html += f"                    <li>{safe_desc}</li>\n"

                html += """                </ul>
            </div>
"""
            html += "        </section>\n"

        # Skills section - use categorized skills
        if resume_data.get("skills"):
            skills_data = await self._deduplicate_and_categorize_skills(
                resume_data["skills"]
            )

            if skills_data:
                html += """
        <section class="resume-section">
            <h2 class="section-title">Skills</h2>
            <div class="skills-grid">
"""
                # Define display order
                category_order = [
                    'Languages', 'Frontend', 'Backend', 'Databases',
                    'Cloud & DevOps', 'Tools', 'Other'
                ]

                for category in category_order:
                    if category in skills_data and skills_data[category]:
                        skills_list = ", ".join(skills_data[category])
                        html += f"""
                <div class="skill-category">
                    <span class="skill-category-name">{category}</span>
                    <span class="skill-inline">{skills_list}</span>
                </div>
"""
                html += """            </div>
        </section>
"""

        # Projects section (sorted: resume projects first, GitHub-sourced last)
        if resume_data.get("projects"):
            html += """
        <section class="resume-section">
            <h2 class="section-title">Projects</h2>
"""
            sorted_projects = self._sort_projects(resume_data["projects"])
            for project in sorted_projects:
                name = project.get('name', '')
                description = project.get('description', '')
                technologies = project.get('technologies', [])
                link = project.get('link', '') or project.get('url', '') or project.get('github_url', '')

                html += f"""
            <div class="project-item">
                <div class="project-header">
                    <span class="project-name">{name}</span>
"""
                if link:
                    display_link = link.replace("https://", "").replace("http://", "")
                    html += f'                    <a href="{link}" class="project-link">{display_link}</a>\n'

                html += f"""                </div>
                <p class="project-description">{description}</p>
"""
                if technologies:
                    tech_str = ", ".join(technologies)
                    html += f"""                <p class="project-tech"><strong>Technologies:</strong> {tech_str}</p>
"""
                html += """            </div>
"""
            html += "        </section>\n"

        # Education section
        if resume_data.get("education"):
            html += """
        <section class="resume-section">
            <h2 class="section-title">Education</h2>
"""
            for edu in resume_data["education"]:
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                year = edu.get('graduation_year', '') or edu.get('year', '')
                gpa = edu.get('gpa', '')

                html += f"""
            <div class="education-item">
                <div class="education-header">
                    <div>
                        <span class="education-degree">{degree}</span>
                        <span class="education-institution"> - {institution}</span>
                    </div>
                    <span class="education-year">{year}</span>
                </div>
"""
                if gpa:
                    html += f"""                <p class="education-details">GPA: {gpa}</p>
"""
                html += """            </div>
"""
            html += "        </section>\n"

        # Certifications section (if present)
        if resume_data.get("certifications"):
            html += """
        <section class="resume-section">
            <h2 class="section-title">Certifications</h2>
"""
            for cert in resume_data["certifications"]:
                name = cert.get('name', '') if isinstance(cert, dict) else cert
                issuer = cert.get('issuer', '') if isinstance(cert, dict) else ''
                date = cert.get('date', '') if isinstance(cert, dict) else ''

                html += f"""
            <div class="certification-item">
                <span class="certification-name">{name}</span>
                <span class="certification-issuer">{issuer} {date}</span>
            </div>
"""
            html += "        </section>\n"

        # Close HTML
        html += """
    </div>
</body>
</html>"""

        return html


export_service = ExportService()